# Create your views here.
#

from django.shortcuts import render, redirect
from django.http import HttpResponse
from .models import ContentTourist
from .process_pdf import html_to_pdf
from django.views.generic import View
from django.template.loader import render_to_string
from docx import Document
from bs4 import BeautifulSoup
import transliterate
def TuristamViews(request):
    return render(
        request,
        'pages/turistam_views.html',
        context={}
    )
def DetailStrTuristam(request, slug_turistampage):
    post = ContentTourist.objects.get(slug=slug_turistampage, is_draft=True)
    return render(
        request,
        'pages/turistam_views_detail.html',
        context={'post':post}
    )


class GeneratePdf(View):
    def get(self, request, slug_turistampage, *args, **kwargs):
        # getting the template
        data = ContentTourist.objects.get(slug=slug_turistampage, is_draft=True)
        open('templates/maket/pdf.html', "w" , encoding='utf-8').write(render_to_string('maket/result.html', {'data': data}))
        pdf = html_to_pdf('maket/pdf.html')
        # rendering the template
        return HttpResponse(pdf, content_type='application/pdf')

def generate_doc(request, slug_turistampage):
    # Создаем новый документ
    document = Document()
    data = ContentTourist.objects.get(slug=slug_turistampage, is_draft=True)

    # Добавляем заголовок в документ
    document.add_heading(data.h1, 0)

    # Удаляем теги HTML и стили из текста и добавляем параграфы в документ
    soup = BeautifulSoup(data.post, 'html.parser')
    content = soup.get_text().split('\n')
    for paragraph in content:
        document.add_paragraph(paragraph)

    # Сохраняем документ в буфер и отправляем его в ответе
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    name = 'attachment; filename="' + transliterate.translit(data.h1, 'ru', reversed=True) + '.docx"'
    response['Content-Disposition'] = str(name)
    document.save(response)
    return response